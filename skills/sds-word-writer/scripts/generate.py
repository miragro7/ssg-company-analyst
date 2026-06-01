#!/usr/bin/env python3
"""
sds-word-writer — SDS 보고서 양식 .docx generator

Usage:
    python generate.py <content.json> <output.docx>

Content JSON schema: see references/content-schema.md
Layout invariants (must stay 1:1 with the Node/JS implementation):
    - Font 바탕체 everywhere (ASCII, EastAsia, hAnsi all 바탕체)
    - Width-scale 95% on every run (장체); per-bullet override via "scale"
    - autoSpaceDE="0" / autoSpaceDN="0" on every paragraph
    - Bold: title, section headings, and explicitly marked bullets
    - Body 14pt (sz=28) black; annotations 9pt (sz=18) #0000FF
    - Single line spacing (line=240)
    - Paragraph spacing.after = 240 (12pt) uniformly for body bullets.
    - Long body bullets are pre-wrapped at word boundaries into multiple
      paragraphs (split_bullet_text) so the 12pt after-spacing applies
      between wrapped lines too.
    - Annotations render as a floating 9pt blue DrawingML text box
      anchored under the first character of the anchor word. No background,
      no border, wrap = in front of text.
    - Colon-based hanging indent: when a bullet contains " : ", wrapped text
      starts at the character after the colon for structural readability.
    - Tables are indented to start at the same x-position as bullet text.
    - Table header row renders bold + centered (benchmark style).
    - Closing "- 이  상 -" is RIGHT-aligned, after=180
"""

from __future__ import annotations

import json
import os
import struct
import subprocess
import sys
import zipfile
import io
import re
from pathlib import Path

# ---------- Runtime dependency bootstrap ----------------------
SKILL_ROOT = Path(__file__).resolve().parent.parent


def _ensure_deps():
    """Install python-docx if missing (first-run bootstrap for ChatGPT sandbox)."""
    try:
        import docx  # noqa: F401
        return
    except ImportError:
        pass
    req = SKILL_ROOT / 'requirements.txt'
    sys.stderr.write('[sds-word-writer] installing python dependencies (python-docx)...\n')
    cmd = [sys.executable, '-m', 'pip', 'install', '--quiet', '--disable-pip-version-check']
    if req.exists():
        cmd += ['-r', str(req)]
    else:
        cmd += ['python-docx']
    subprocess.check_call(cmd)


_ensure_deps()

from docx import Document  # noqa: E402
from docx.shared import Twips, Pt, Emu, RGBColor  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from lxml import etree  # noqa: E402


# ---------- Style constants -----------------------------------
FONT_NAMES = {'ascii': '바탕체', 'eastAsia': '바탕체', 'hAnsi': '바탕체'}
WIDTH_PCT = 95
BLUE = '0000FF'
BLACK = '000000'
SIZE_TITLE = 40   # 20pt (half-points)
SIZE_BODY = 28    # 14pt
SIZE_ANNOT = 18   # 9pt
SIZE_TABLE = 24   # 12pt
SIZE_APPEND = 28  # 14pt

# Character-width presets (DXA = 1/20 pt). Body 14pt 바탕체 + 장체 95%.
BODY_PT = 14
FULLWIDTH_DXA = round(BODY_PT * 20 * 1.00 * (WIDTH_PCT / 100))  # ~266
HALFWIDTH_DXA = round(BODY_PT * 20 * 0.50 * (WIDTH_PCT / 100))  # ~133

# Annotation 9pt 바탕체 + 장체 95%
ANNOT_PT = 9
ANNOT_FULLWIDTH_DXA = round(ANNOT_PT * 20 * 1.00 * (WIDTH_PCT / 100))  # ~171
ANNOT_HALFWIDTH_DXA = round(ANNOT_PT * 20 * 0.50 * (WIDTH_PCT / 100))  # ~86

AFTER = {
    'title': 240,
    'date': 240,
    'section_heading': 240,
    'body': 240,
    'body_wrap': 180,
    'last_in_section': 180,
    'section_separator': 240,
    'annotation': 240,
    'closing': 180,
    'append_label': 240,
}

PAGE = {
    'width': 11906, 'height': 16838,
    'top': 1134, 'right': 1134, 'bottom': 1134, 'left': 1134,
    'header': 851, 'footer': 510,
}
PAGE_CONTENT_WIDTH = PAGE['width'] - PAGE['left'] - PAGE['right']
TWIPS_PER_INCH = 1440
PX_PER_INCH = 96

FLOAT_ANNOT_Y = -240
WRAP_SAFETY_MARGIN = 0.99

# EMU / DXA conversion for DrawingML (1 inch = 914400 EMU = 1440 DXA)
EMU_PER_DXA = 635
ANNOT_TEXTBOX_HEIGHT_DXA = 260      # single-line 9pt text box height (a bit > line=200)
ANNOT_TEXTBOX_Y_OFFSET_DXA = -240   # pull the frame up so it overlays the anchor bullet line

_DRAWING_ID_COUNTER = [1000]


def _next_drawing_id() -> int:
    _DRAWING_ID_COUNTER[0] += 1
    return _DRAWING_ID_COUNTER[0]


def _emu_from_dxa(dxa: int) -> int:
    return int(round(dxa * EMU_PER_DXA))

PREFIX = {
    'dash':  '  - ',
    'dot':   '    \u00B7',
    'wrap':  '      ',
    'note':  '  \u203B ',
    'arrow': '    \u2192 ',
}
CIRCLED = ['\u2460', '\u2461', '\u2462', '\u2463', '\u2464',
           '\u2465', '\u2466', '\u2467', '\u2468', '\u2469']


# ---------- Low-level XML helpers -----------------------------
def _qn(name: str) -> str:
    """Qualified name helper: '_qn("val")' => w:val."""
    return qn('w:' + name) if ':' not in name else qn(name)


def _make(tag: str, **attrs) -> OxmlElement:
    """Create an OxmlElement with the given attrs. Keys default to w:<key>.

    Tag accepts either 'rFonts' (auto-prefixed to 'w:rFonts') or 'w:rFonts'.
    """
    if ':' not in tag:
        tag = 'w:' + tag
    el = OxmlElement(tag)
    for k, v in attrs.items():
        el.set(_qn(k), str(v))
    return el


def _set_attr(el, name: str, value) -> None:
    el.set(_qn(name), str(value))


# ---------- Unicode / measurement -----------------------------
def is_fullwidth(code: int) -> bool:
    """Matches the JS isFullwidth() — CJK / Hangul / wide punctuation set."""
    return (
        (0x1100 <= code <= 0x115F) or
        (0x2E80 <= code <= 0x303E) or
        (0x3041 <= code <= 0x33FF) or
        (0x3400 <= code <= 0x4DBF) or
        (0x4E00 <= code <= 0x9FFF) or
        (0xA000 <= code <= 0xA4CF) or
        (0xAC00 <= code <= 0xD7A3) or
        (0xF900 <= code <= 0xFAFF) or
        (0xFE30 <= code <= 0xFE4F) or
        (0xFF00 <= code <= 0xFF60) or
        (0xFFE0 <= code <= 0xFFE6) or
        code == 0x00B7 or
        code == 0x203B or
        code == 0x223C or
        (0x2014 <= code <= 0x2015) or
        (0x2190 <= code <= 0x21FF) or
        (0x25A0 <= code <= 0x25FF) or
        (0x2460 <= code <= 0x24FF)
    )


def measure_dxa(s: str, scale: int = WIDTH_PCT) -> int:
    """Measure string width in DXA. `scale` is character-width percent (장평).

    Fast path when scale == WIDTH_PCT uses precomputed constants; otherwise
    recomputes per-char widths for the requested scale.
    """
    if scale == WIDTH_PCT:
        dxa = 0
        for ch in s or '':
            dxa += FULLWIDTH_DXA if is_fullwidth(ord(ch)) else HALFWIDTH_DXA
        return dxa
    full = round(BODY_PT * 20 * 1.00 * (scale / 100))
    half = round(BODY_PT * 20 * 0.50 * (scale / 100))
    dxa = 0
    for ch in s or '':
        dxa += full if is_fullwidth(ord(ch)) else half
    return dxa


def measure_annot_dxa(s: str) -> int:
    dxa = 0
    for ch in s or '':
        dxa += ANNOT_FULLWIDTH_DXA if is_fullwidth(ord(ch)) else ANNOT_HALFWIDTH_DXA
    return dxa


def compute_anchor_indent_dxa(prefix: str, text: str, anchor: str | None) -> int:
    if not anchor:
        return 0
    idx = (text or '').find(anchor)
    if idx < 0:
        return 0
    return measure_dxa((prefix or '') + (text or '')[:idx])


def hanging_indent_for(prefix: str, text: str):
    """Returns dict {'left': DXA, 'hanging': DXA} or None.

    Colon-based indent: when text contains ' : ', hang at character-after-colon
    (only if that point lies within 55% of page content width).
    """
    if text:
        idx = text.find(' : ')
        if idx >= 0:
            colon_end = idx + 3
            full_prefix = prefix + text[:colon_end]
            dxa = measure_dxa(full_prefix)
            if dxa < PAGE_CONTENT_WIDTH * 0.55:
                return {'left': dxa, 'hanging': dxa}
    dxa = measure_dxa(prefix)
    if dxa <= 0:
        return None
    return {'left': dxa, 'hanging': dxa}


def normalize_annotation(annotation):
    if not annotation:
        return None
    if isinstance(annotation, str):
        return {'text': annotation, 'anchor': None}
    if isinstance(annotation, dict):
        return {
            'text': annotation.get('text', ''),
            'anchor': annotation.get('anchor') or None,
        }
    return None


def split_bullet_text(prefix: str, text: str, first_line_capacity: int, continuation_capacity: int,
                      prefix_scale: int = WIDTH_PCT, text_scale: int = WIDTH_PCT):
    """Pre-wrap text at word boundaries (see JS splitBulletText).

    Returns list of line strings; each line becomes its own paragraph so the
    12pt after-spacing applies between wrapped lines.

    `prefix_scale` / `text_scale` let the caller measure prefix (bullet mark)
    and sentence body at different character-width percents — used by the
    auto widow-fix pass which shrinks only the body's 장평.
    """
    src = text or ''
    prefix_dxa = measure_dxa(prefix or '', prefix_scale)
    if prefix_dxa + measure_dxa(src, text_scale) <= first_line_capacity:
        return [src]
    words = [w for w in re.split(r' +', src) if w]
    if len(words) <= 1:
        return [src]

    lines = []
    cur: list[str] = []
    cur_dxa = 0
    cap = first_line_capacity - prefix_dxa

    def flush():
        nonlocal cur, cur_dxa, cap
        lines.append(' '.join(cur))
        cur = []
        cur_dxa = 0
        cap = continuation_capacity

    space_dxa = round(BODY_PT * 20 * 0.50 * (text_scale / 100))
    for word in words:
        w_dxa = measure_dxa(word, text_scale)
        delta = (space_dxa + w_dxa) if cur else w_dxa
        if cur and cur_dxa + delta > cap:
            flush()
            cur.append(word)
            cur_dxa = w_dxa
        else:
            cur.append(word)
            cur_dxa += delta
    if cur:
        flush()
    return lines


def _last_line_is_widow(lines) -> bool:
    """Return True if the bullet wrapped to ≥2 lines and the final line has
    only 1–2 whitespace-separated words (typical widow pattern).
    """
    if not lines or len(lines) < 2:
        return False
    last = (lines[-1] or '').strip()
    if not last:
        return False
    words = [w for w in last.split() if w]
    return len(words) <= 2


# ---------- Image header parsing (no Pillow needed) -----------
def read_png_size(buf: bytes):
    if len(buf) < 24 or buf[1:4] != b'PNG':
        return None
    width = struct.unpack('>I', buf[16:20])[0]
    height = struct.unpack('>I', buf[20:24])[0]
    return {'width': width, 'height': height}


def read_jpeg_size(buf: bytes):
    if len(buf) < 4 or buf[0] != 0xFF or buf[1] != 0xD8:
        return None
    offset = 2
    n = len(buf)
    while offset + 9 < n:
        if buf[offset] != 0xFF:
            offset += 1
            continue
        marker = buf[offset + 1]
        if marker in (0xD8, 0xD9):
            offset += 2
            continue
        length = struct.unpack('>H', buf[offset + 2:offset + 4])[0]
        if length < 2 or offset + 2 + length > n:
            return None
        is_sof = (
            0xC0 <= marker <= 0xC3 or
            0xC5 <= marker <= 0xC7 or
            0xC9 <= marker <= 0xCB or
            0xCD <= marker <= 0xCF
        )
        if is_sof:
            height = struct.unpack('>H', buf[offset + 5:offset + 7])[0]
            width = struct.unpack('>H', buf[offset + 7:offset + 9])[0]
            return {'width': width, 'height': height}
        offset += 2 + length
    return None


def read_image_size(path: str, buf: bytes):
    ext = os.path.splitext(path)[1].lower()
    if ext == '.png':
        return read_png_size(buf)
    if ext in ('.jpg', '.jpeg'):
        return read_jpeg_size(buf)
    return None


def twips_to_px(twips: int) -> int:
    return round((twips / TWIPS_PER_INCH) * PX_PER_INCH)


def image_dimensions(bullet: dict, path: str, buf: bytes):
    intrinsic = read_image_size(path, buf)
    width = bullet.get('width') or bullet.get('widthPx') or (intrinsic and intrinsic.get('width'))
    height = bullet.get('height') or bullet.get('heightPx') or (intrinsic and intrinsic.get('height'))
    if width and not height and intrinsic and intrinsic.get('width') and intrinsic.get('height'):
        height = round(width * intrinsic['height'] / intrinsic['width'])
    elif height and not width and intrinsic and intrinsic.get('width') and intrinsic.get('height'):
        width = round(height * intrinsic['width'] / intrinsic['height'])
    elif not width or not height:
        width = (intrinsic and intrinsic.get('width')) or 600
        height = (intrinsic and intrinsic.get('height')) or 300
    max_width = max(120, round(twips_to_px(PAGE_CONTENT_WIDTH) * (bullet.get('maxWidthPct') or 0.92)))
    if width > max_width:
        scale = max_width / width
        width = round(width * scale)
        height = round(height * scale)
    return width, height


def resolve_image_path(image_path: str) -> str:
    if not image_path:
        raise ValueError('Image bullet requires "path".')
    return image_path if os.path.isabs(image_path) else os.path.abspath(image_path)


# ---------- Run / paragraph primitives ------------------------
def _set_run_font(rpr, font_names=FONT_NAMES):
    fonts = _make('rFonts',
                  ascii=font_names['ascii'],
                  eastAsia=font_names['eastAsia'],
                  hAnsi=font_names['hAnsi'])
    # also set cs (complex script) to match
    fonts.set(_qn('cs'), font_names['ascii'])
    rpr.append(fonts)


def _set_run_props(r_element, *, size=SIZE_BODY, color=BLACK, bold=False,
                   underline=False, scale=WIDTH_PCT, character_spacing=None):
    """Attach rPr to a w:r element with SDS-standard run properties."""
    rpr = r_element.find(_qn('rPr'))
    if rpr is None:
        rpr = _make('rPr')
        r_element.insert(0, rpr)
    else:
        # clear to avoid duplicates
        for child in list(rpr):
            rpr.remove(child)

    _set_run_font(rpr)
    if bold:
        rpr.append(_make('b'))
        rpr.append(_make('bCs'))
    if underline:
        rpr.append(_make('u', val='single'))
    rpr.append(_make('color', val=color))
    rpr.append(_make('sz', val=size))
    rpr.append(_make('szCs', val=size))
    rpr.append(_make('w', val=scale))
    if character_spacing is not None:
        rpr.append(_make('spacing', val=character_spacing))


def make_run(text: str, *, size=SIZE_BODY, color=BLACK, bold=False,
             underline=False, scale=WIDTH_PCT, character_spacing=None) -> OxmlElement:
    r = _make('r')
    _set_run_props(r, size=size, color=color, bold=bold, underline=underline,
                   scale=scale, character_spacing=character_spacing)
    t = _make('t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text or ''
    r.append(t)
    return r


def make_annot_run(text: str, prefix_star: bool = True) -> OxmlElement:
    t = f'* {text}' if (prefix_star and not (text or '').startswith('*')) else text
    return make_run(t, size=SIZE_ANNOT, color=BLUE, scale=WIDTH_PCT)


_ALIGN = {
    'left': 'left',
    'right': 'right',
    'center': 'center',
    'both': 'both',
}


def _apply_p_properties(p_element, *, alignment=None, line=240, line_rule='auto',
                        before=None, after=None, indent=None, frame=None,
                        keep_bold_empty=False):
    """Append a pPr to a w:p element with the given properties."""
    pPr = p_element.find(_qn('pPr'))
    if pPr is None:
        pPr = _make('pPr')
        p_element.insert(0, pPr)
    # We do NOT pre-inject autoSpaceDE/DN here — postProcessAutoSpace adds them.
    # Frame first if present
    if frame is not None:
        pPr.append(frame)
    if alignment in _ALIGN:
        pPr.append(_make('jc', val=_ALIGN[alignment]))
    spacing_attrs = {}
    if line is not None:
        spacing_attrs['line'] = line
        spacing_attrs['lineRule'] = line_rule
    if before is not None:
        spacing_attrs['before'] = before
    if after is not None:
        spacing_attrs['after'] = after
    if spacing_attrs:
        pPr.append(_make('spacing', **spacing_attrs))
    if indent is not None:
        ind_attrs = {}
        if 'left' in indent and indent['left']:
            ind_attrs['left'] = indent['left']
        if 'hanging' in indent and indent['hanging']:
            ind_attrs['hanging'] = indent['hanging']
        if 'firstLine' in indent and indent['firstLine']:
            ind_attrs['firstLine'] = indent['firstLine']
        if ind_attrs:
            pPr.append(_make('ind', **ind_attrs))
    if keep_bold_empty:
        # rPr on pPr affects paragraph mark only — used for bold empty separators.
        rpr = _make('rPr')
        rpr.append(_make('b'))
        pPr.append(rpr)


def make_paragraph(runs, *, alignment=None, line=240, line_rule='auto',
                   before=None, after=None, indent=None, frame=None,
                   keep_bold_empty=False) -> OxmlElement:
    p = _make('p')
    _apply_p_properties(p, alignment=alignment, line=line, line_rule=line_rule,
                        before=before, after=after, indent=indent, frame=frame,
                        keep_bold_empty=keep_bold_empty)
    for r in (runs or []):
        p.append(r)
    return p


def body_paragraph(text: str, *, alignment=None, after=AFTER['body'],
                   size=SIZE_BODY, color=BLACK, indent=None, bold=False,
                   underline=False, scale=WIDTH_PCT) -> OxmlElement:
    r = make_run(text, size=size, color=color, bold=bold, underline=underline, scale=scale)
    return make_paragraph([r], alignment=alignment, line=240, line_rule='auto',
                          after=after, indent=indent)


# ---------- Annotation paragraphs -----------------------------
_NS_W  = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
_NS_WP = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
_NS_A  = 'http://schemas.openxmlformats.org/drawingml/2006/main'
_NS_WPS = 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'


def _xml_escape(s: str) -> str:
    return (s or '').replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')


def _build_annotation_drawing(text: str, *, indent_dxa: int):
    """Build a w:drawing element containing a DrawingML text box (텍스트 상자).

    - 배경 없음 (a:noFill), 테두리 없음 (a:ln/a:noFill)
    - 텍스트 줄 바꿈: 텍스트 앞 (wp:wrapNone + behindDoc="0")
    - 바탕체 9pt #0000FF, 장체 95%
    """
    ann_text = f'* {text}' if not (text or '').startswith('*') else text
    text_width_dxa = measure_annot_dxa(ann_text)
    padding_dxa = 200
    max_width_dxa = max(PAGE_CONTENT_WIDTH - max(indent_dxa, 0), 2000)
    frame_width_dxa = min(text_width_dxa + padding_dxa, max_width_dxa)

    width_emu = _emu_from_dxa(frame_width_dxa)
    height_emu = _emu_from_dxa(ANNOT_TEXTBOX_HEIGHT_DXA)
    x_emu = _emu_from_dxa(max(indent_dxa, 0))
    y_emu = _emu_from_dxa(ANNOT_TEXTBOX_Y_OFFSET_DXA)

    drawing_id = _next_drawing_id()
    escaped = _xml_escape(ann_text)

    xml = f'''<w:drawing xmlns:w="{_NS_W}" xmlns:wp="{_NS_WP}" xmlns:a="{_NS_A}" xmlns:wps="{_NS_WPS}">
  <wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0"
             relativeHeight="{251659000 + drawing_id}" behindDoc="0"
             locked="0" layoutInCell="1" allowOverlap="1">
    <wp:simplePos x="0" y="0"/>
    <wp:positionH relativeFrom="column"><wp:posOffset>{x_emu}</wp:posOffset></wp:positionH>
    <wp:positionV relativeFrom="paragraph"><wp:posOffset>{y_emu}</wp:posOffset></wp:positionV>
    <wp:extent cx="{width_emu}" cy="{height_emu}"/>
    <wp:effectExtent l="0" t="0" r="0" b="0"/>
    <wp:wrapNone/>
    <wp:docPr id="{drawing_id}" name="Annotation {drawing_id}"/>
    <wp:cNvGraphicFramePr/>
    <a:graphic>
      <a:graphicData uri="{_NS_WPS}">
        <wps:wsp>
          <wps:cNvSpPr txBox="1"/>
          <wps:spPr>
            <a:xfrm>
              <a:off x="0" y="0"/>
              <a:ext cx="{width_emu}" cy="{height_emu}"/>
            </a:xfrm>
            <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
            <a:noFill/>
            <a:ln><a:noFill/></a:ln>
          </wps:spPr>
          <wps:txbx>
            <w:txbxContent>
              <w:p>
                <w:pPr>
                  <w:spacing before="0" after="0" line="200" lineRule="exact"/>
                  <w:rPr>
                    <w:rFonts w:ascii="바탕체" w:eastAsia="바탕체" w:hAnsi="바탕체" w:cs="바탕체"/>
                    <w:color w:val="{BLUE}"/>
                    <w:sz w:val="{SIZE_ANNOT}"/>
                    <w:szCs w:val="{SIZE_ANNOT}"/>
                    <w:w w:val="{WIDTH_PCT}"/>
                  </w:rPr>
                </w:pPr>
                <w:r>
                  <w:rPr>
                    <w:rFonts w:ascii="바탕체" w:eastAsia="바탕체" w:hAnsi="바탕체" w:cs="바탕체"/>
                    <w:color w:val="{BLUE}"/>
                    <w:sz w:val="{SIZE_ANNOT}"/>
                    <w:szCs w:val="{SIZE_ANNOT}"/>
                    <w:w w:val="{WIDTH_PCT}"/>
                  </w:rPr>
                  <w:t xml:space="preserve">{escaped}</w:t>
                </w:r>
              </w:p>
            </w:txbxContent>
          </wps:txbx>
          <wps:bodyPr rot="0" spcFirstLastPara="0" vertOverflow="overflow"
                      horzOverflow="overflow" wrap="square"
                      lIns="0" tIns="0" rIns="0" bIns="0"
                      anchor="t" anchorCtr="0">
            <a:noAutofit/>
          </wps:bodyPr>
        </wps:wsp>
      </a:graphicData>
    </a:graphic>
  </wp:anchor>
</w:drawing>'''
    return etree.fromstring(xml)


def annotation_textbox(text: str, *, indent_dxa=0, after=0) -> OxmlElement:
    """Host paragraph carrying a floating DrawingML text box annotation.

    The host paragraph is intentionally near-zero height (line=20 exact,
    before/after controlled by caller) — the annotation itself is anchored
    to this paragraph and overlays the preceding bullet line in "텍스트 앞"
    wrap mode.
    """
    drawing = _build_annotation_drawing(text, indent_dxa=indent_dxa)
    r = _make('r')
    r.append(drawing)
    return make_paragraph([r], line=20, line_rule='exact', before=0, after=after)


def annotation_paragraph(text: str, *, indent_dxa=0, after=AFTER['annotation']) -> OxmlElement:
    return annotation_textbox(text, indent_dxa=indent_dxa, after=after)


# ---------- Image paragraph ------------------------------------
def _emu_from_px(px: int) -> int:
    # 1 inch = 96 px = 914400 EMU
    return int(round(px * 914400 / 96))


def render_image_bullet(doc: Document, bullet: dict, is_last: bool = False):
    """Appends an image paragraph (and optional caption) to the document body.

    Note: this uses python-docx add_picture() which creates a paragraph in the
    current body position with an inline shape, then we tweak the paragraph.
    """
    file_path = resolve_image_path(bullet.get('path', ''))
    with open(file_path, 'rb') as f:
        data = f.read()
    width_px, height_px = image_dimensions(bullet, file_path, data)

    # Use add_picture helper: it appends a paragraph with an inline drawing.
    p_runcarrier = doc.add_paragraph()
    run = p_runcarrier.add_run()
    run.add_picture(io.BytesIO(data), width=Emu(_emu_from_px(width_px)),
                    height=Emu(_emu_from_px(height_px)))
    # Alignment
    p_runcarrier.alignment = (WD_ALIGN_PARAGRAPH.LEFT
                              if bullet.get('alignment') == 'left'
                              else WD_ALIGN_PARAGRAPH.CENTER)
    # Spacing
    after_val = 0 if bullet.get('caption') else (AFTER['last_in_section'] if is_last else AFTER['body'])
    _apply_p_properties(p_runcarrier._p, line=240, line_rule='auto', after=after_val)

    if bullet.get('caption'):
        cap_p = body_paragraph(f"  \u203B {bullet['caption']}",
                               after=(AFTER['last_in_section'] if is_last else AFTER['body']),
                               size=SIZE_TABLE)
        _append_block(doc, cap_p)


# ---------- Body insertion helper -----------------------------
def _append_block(doc: Document, element: OxmlElement) -> None:
    """Append a block element (w:p or w:tbl) to the document body, before sectPr."""
    body = doc.element.body
    sect_pr = body.find(qn('w:sectPr'))
    if sect_pr is not None:
        sect_pr.addprevious(element)
    else:
        body.append(element)


def _append_blocks(doc: Document, elements) -> None:
    for el in elements:
        _append_block(doc, el)


# ---------- Table rendering -----------------------------------
NUMERIC_CELL_REGEX = re.compile(
    r'^[()\-+]?[\d,]+(\.\d+)?\s*(%p?|%|원|₩|달러|\$|B\$|M\$|억|만|천|조|pp|p|배|명|건|시간|분|일|개|kg|g|\s)*\)?$'
)
LONG_NUMERIC_MARKER = re.compile(r'[,.]|\$|₩|달러|B\$|M\$|억|만|천|조|원')
# Columns whose header names this pattern are money/scale quantities — the
# SDS convention is right-align even when the sampled values happen to be
# all short (no thousands-separator) so a column with mixed "820" and "1,200"
# reads consistently. Applied before the short-comparable/center fallback.
MONEY_HEADER_RE = re.compile(
    r'금액|매출|영업이익|순이익|당기순이익|영업손실|손실|비용|수익|세액|'
    r'원$|억원|만원|백만원|천만원|달러|amount|price|revenue|cost|profit|sales',
    re.IGNORECASE,
)


def _is_money_header(text) -> bool:
    if text is None:
        return False
    s = str(text).strip()
    if not s:
        return False
    return bool(MONEY_HEADER_RE.search(s))

CELL_H_PADDING_DXA = 200
MIN_COL_WIDTH_DXA = 900


def is_numeric_cell(v) -> bool:
    if v is None:
        return False
    s = str(v).strip()
    if s == '' or s == '-':
        return False
    return bool(NUMERIC_CELL_REGEX.match(s))


def is_fixed_length_column(values) -> bool:
    non_empty = [v for v in values if v is not None and str(v).strip() != '']
    if len(non_empty) < 2:
        return False
    length = len(str(non_empty[0]).strip())
    if length == 0 or length > 8:
        return False
    return all(len(str(v).strip()) == length for v in non_empty)


def is_short_comparable_numeric_column(values) -> bool:
    non_empty = [v for v in values if v is not None and str(v).strip() != '']
    if len(non_empty) < 2:
        return False
    if not all(is_numeric_cell(v) for v in non_empty):
        return False
    for v in non_empty:
        s = str(v).strip()
        if len(s) > 5 or LONG_NUMERIC_MARKER.search(s):
            return False
    return True


def infer_column_alignments(cols, rows):
    result = []
    for col_idx, header in enumerate(cols):
        col_vals = [r[col_idx] if col_idx < len(r) else None for r in rows]
        non_empty = [v for v in col_vals if v is not None and str(v).strip() != '']
        if not non_empty:
            result.append('left')
            continue
        all_num = all(is_numeric_cell(v) for v in non_empty)
        # Money/scale header with numeric body → right, even when values are
        # short digits without thousands separators.
        if all_num and _is_money_header(header):
            result.append('right')
            continue
        if all_num and any(LONG_NUMERIC_MARKER.search(str(v).strip()) for v in non_empty):
            result.append('right')
            continue
        if is_fixed_length_column(col_vals):
            result.append('center')
            continue
        if is_short_comparable_numeric_column(col_vals):
            result.append('center')
            continue
        if all_num:
            result.append('right')
            continue
        result.append('left')
    return result


def column_content_dxa(header, body_values, font_pt):
    scale = font_pt / 14
    cells = [str(header or '')] + [str(v or '') for v in body_values]
    max_dxa = max((measure_dxa(c) for c in cells), default=0)
    return round(max_dxa * scale) + CELL_H_PADDING_DXA


def is_compressible_column(body_values) -> bool:
    non_empty = [v for v in body_values if v is not None and str(v).strip() != '']
    if not non_empty:
        return True
    if all(is_numeric_cell(v) for v in non_empty):
        return False
    if all(len(str(v).strip()) <= 6 for v in non_empty):
        return False
    return True


def compute_auto_column_widths(cols, rows, total_width, font_pt):
    n = len(cols)
    if n == 0:
        return []
    naturals = [
        max(column_content_dxa(h, [r[i] if i < len(r) else None for r in rows], font_pt),
            MIN_COL_WIDTH_DXA)
        for i, h in enumerate(cols)
    ]
    total = sum(naturals)
    if total >= total_width:
        return [int(w * total_width / total) for w in naturals]
    extra = total_width - total
    compressible = [is_compressible_column([r[i] if i < len(r) else None for r in rows])
                    for i in range(n)]
    flex_weight_total = sum(naturals[i] for i in range(n) if compressible[i])
    if flex_weight_total == 0:
        per = extra // n
        return [w + per for w in naturals]
    return [
        w + int(extra * w / flex_weight_total) if compressible[i] else w
        for i, w in enumerate(naturals)
    ]


def resolve_column_widths(b, cols, rows, total_width, font_pt):
    cw = b.get('columnWidths')
    if isinstance(cw, list) and len(cw) == len(cols):
        try:
            weights = [float(w) if (w is not None and float(w) > 0) else 1.0 for w in cw]
        except (TypeError, ValueError):
            weights = [1.0] * len(cols)
        s = sum(weights)
        return [int(total_width * w / s) for w in weights]
    return compute_auto_column_widths(cols, rows, total_width, font_pt)


def _alignment_from_string(s):
    if s == 'right':
        return 'right'
    if s == 'center':
        return 'center'
    return 'left'


def _clamp_table_font_pt(pt):
    try:
        n = float(pt)
    except (TypeError, ValueError):
        return 12
    if n <= 0:
        return 12
    return max(9, min(14, n))


def _border_el(name: str):
    return _make(name, val='single', sz=4, color='000000', space=0)


def _make_cell(text, *, is_header: bool, col_width: int, font_size_half_pt: int,
               alignment: str) -> OxmlElement:
    tc = _make('tc')
    tcPr = _make('tcPr')
    tcPr.append(_make('tcW', w=col_width, type='dxa'))
    borders = _make('tcBorders')
    for b in ('top', 'left', 'bottom', 'right'):
        borders.append(_border_el(b))
    tcPr.append(borders)
    tcPr.append(_make('vAlign', val='center'))
    mar = _make('tcMar')
    mar.append(_make('top', w=60, type='dxa'))
    mar.append(_make('bottom', w=60, type='dxa'))
    mar.append(_make('left', w=100, type='dxa'))
    mar.append(_make('right', w=100, type='dxa'))
    tcPr.append(mar)
    tc.append(tcPr)
    run = make_run(str(text or ''), size=font_size_half_pt, bold=is_header)
    p = make_paragraph([run],
                       alignment=('center' if is_header else _alignment_from_string(alignment)),
                       line=240, line_rule='auto', after=0)
    tc.append(p)
    return tc


# ---------- Table merge (cells-mode) helpers ------------------
def _apply_keep_next(p_element) -> None:
    """Add w:keepNext to a paragraph so it stays on the same page as the next block."""
    pPr = p_element.find(_qn('pPr'))
    if pPr is None:
        pPr = _make('pPr')
        p_element.insert(0, pPr)
    if pPr.find(_qn('keepNext')) is None:
        pPr.insert(0, _make('keepNext'))


def _apply_table_keep_together(tbl) -> None:
    """Keep an entire table on one page when possible.

    Two-layer strategy:
      1. `w:cantSplit` on every row — prevents a single row from splitting
         mid-cell across a page boundary.
      2. `w:keepNext` on every paragraph inside every row EXCEPT the last
         row — forces Word to keep each row attached to the next one, so
         the whole table moves to the next page together when it doesn't
         fit. cantSplit alone is not enough; without keepNext Word will
         happily break between rows.

    We skip the final row so Word allows the block AFTER the table (the
    trailing spacer paragraph) to flow normally, and so we add one fewer
    ■ indicator per table in the editing view.
    """
    rows = tbl.findall(_qn('tr'))
    for idx, tr in enumerate(rows):
        trPr = tr.find(_qn('trPr'))
        if trPr is None:
            trPr = _make('trPr')
            tr.insert(0, trPr)
        if trPr.find(_qn('cantSplit')) is None:
            trPr.append(_make('cantSplit'))
        if idx == len(rows) - 1:
            continue
        for tc in tr.findall(_qn('tc')):
            for p in tc.findall(_qn('p')):
                _apply_keep_next(p)


def _infer_one_column_alignment(values, header_texts=()) -> str:
    """Per-column alignment for a single body column (cells-mode).

    Mirrors the logic of `infer_column_alignments` but runs against one column
    at a time. Uses aggregate signals across the column so every cell in the
    column ends up with the same alignment — e.g. a 금액 column with both
    `820` and `1,200` is classified as right-aligned for all rows, instead of
    mixing center/right per-cell.

    `header_texts` are any header-row origin texts that cover this column. If
    any of them matches a money/scale keyword (`금액`, `매출`, …) we force
    right alignment even when the sampled body values are all short digits.
    """
    non_empty = [v for v in values if v is not None and str(v).strip() != '']
    if not non_empty:
        return 'left'
    all_num = all(is_numeric_cell(v) for v in non_empty)
    if all_num and any(_is_money_header(h) for h in header_texts):
        return 'right'
    if all_num and any(LONG_NUMERIC_MARKER.search(str(v).strip()) for v in non_empty):
        return 'right'
    if is_fixed_length_column(values):
        return 'center'
    if is_short_comparable_numeric_column(values):
        return 'center'
    if all_num:
        return 'right'
    return 'left'


def _make_merged_cell(text, *, col_widths, colspan, is_header, bold,
                      alignment, font_size_half_pt, vmerge=None):
    """Build a w:tc with optional gridSpan / vMerge."""
    tc = _make('tc')
    tcPr = _make('tcPr')
    total_w = sum(col_widths) if col_widths else 900
    tcPr.append(_make('tcW', w=total_w, type='dxa'))
    if colspan > 1:
        tcPr.append(_make('gridSpan', val=colspan))
    if vmerge == 'restart':
        tcPr.append(_make('vMerge', val='restart'))
    elif vmerge == 'continue':
        tcPr.append(_make('vMerge'))  # no val = continue
    borders = _make('tcBorders')
    for bn in ('top', 'left', 'bottom', 'right'):
        borders.append(_border_el(bn))
    tcPr.append(borders)
    tcPr.append(_make('vAlign', val='center'))
    mar = _make('tcMar')
    mar.append(_make('top', w=60, type='dxa'))
    mar.append(_make('bottom', w=60, type='dxa'))
    mar.append(_make('left', w=100, type='dxa'))
    mar.append(_make('right', w=100, type='dxa'))
    tcPr.append(mar)
    tc.append(tcPr)
    run = make_run(str(text or ''), size=font_size_half_pt, bold=bold)
    p = make_paragraph([run], alignment=_alignment_from_string(alignment),
                       line=240, line_rule='auto', after=0)
    tc.append(p)
    return tc


def _render_cells_table(b: dict) -> OxmlElement:
    """Render a table from a 2-D `cells` array with rowspan/colspan support.

    cells item forms:
      - str                       → single cell, text only
      - {"text", "rowspan",       → explicit cell descriptor
         "colspan", "align", "bold"}
      - None                      → position covered by a prior merge
    """
    cells = b.get('cells') or []
    header_row_count = max(int(b.get('headerRowCount') or 1), 0)
    font_pt = _clamp_table_font_pt(b.get('fontSize', 12))
    font_size_half_pt = round(font_pt * 2)

    # Determine total logical column count: max over all rows of (sum of
    # colspans for dict cells + 1 for each string/None item).
    n_cols = 0
    for row in cells:
        width = 0
        for item in row or []:
            if isinstance(item, dict):
                width += max(int(item.get('colspan', 1)), 1)
            else:
                width += 1
        if width > n_cols:
            n_cols = width
    if n_cols == 0:
        return _make('tbl')  # empty shell

    # Column widths — always sum exactly to total_width so the table's outer
    # right edge aligns precisely with the page content right margin (which is
    # where a right-aligned caption/unit paragraph above the table ends up).
    table_indent = measure_dxa(PREFIX['dash'])
    total_width = PAGE_CONTENT_WIDTH - table_indent
    cw_in = b.get('columnWidths')
    if isinstance(cw_in, list) and len(cw_in) == n_cols:
        try:
            weights = [max(float(w), 0.01) for w in cw_in]
        except (TypeError, ValueError):
            weights = [1.0] * n_cols
        s = sum(weights) or 1.0
        column_widths = [int(total_width * w / s) for w in weights]
    else:
        base = total_width // n_cols
        column_widths = [base] * n_cols
    # Distribute rounding remainder across first N columns so the sum equals
    # total_width exactly.
    remainder = total_width - sum(column_widths)
    for i in range(max(remainder, 0)):
        column_widths[i % n_cols] += 1

    # Build an origin map + coverage status grid
    n_rows = len(cells)
    status = [[None] * n_cols for _ in range(n_rows)]
    origin_data = {}  # (r, c) -> dict
    # Grid contract: each row is a full 2-D grid — exactly one item per
    # logical column. `null` marks a slot covered by a prior merge;
    # strings/dicts emit origin cells (a dict's colspan consumes N columns).
    for r in range(n_rows):
        row = cells[r] or []
        c = 0
        for item in row:
            if c >= n_cols:
                break
            if item is None:
                c += 1
                continue
            if isinstance(item, dict):
                rs = max(int(item.get('rowspan', 1)), 1)
                cs = max(int(item.get('colspan', 1)), 1)
                text = item.get('text', '')
                align_ov = item.get('align')
                bold_ov = item.get('bold')
            else:
                rs, cs = 1, 1
                text = str(item)
                align_ov = None
                bold_ov = None
            rs = min(rs, n_rows - r)
            cs = min(cs, n_cols - c)
            # If this slot is already covered by an earlier merge, the user's
            # item is in conflict — skip it (advance by cs) and keep going.
            if status[r][c] is not None:
                c += cs
                continue
            status[r][c] = ('ORIGIN', rs, cs)
            origin_data[(r, c)] = {
                'text': text, 'rowspan': rs, 'colspan': cs,
                'align': align_ov, 'bold': bold_ov,
            }
            for cc in range(c + 1, c + cs):
                status[r][cc] = ('HCOVER',)
            for rr in range(r + 1, r + rs):
                status[rr][c] = ('VCOVER_LEFT', cs)
                for cc in range(c + 1, c + cs):
                    status[rr][cc] = ('VCOVER_MID',)
            c += cs

    # Column-level alignment inference for body leaf cells (colspan=1,
    # rowspan=1 only). Matches the simple-mode behavior: a column with any
    # thousands-separated / money-marker value aligns all its cells right, so
    # `820` and `1,200` in the same 금액 column don't split across center/right.
    col_alignments = []
    # Pre-collect header-row origin texts per logical column so the money-
    # header heuristic works across spanned headers too (e.g. a column whose
    # innermost header is "금액" under a "'25.상반기" group header).
    header_texts_by_col = [[] for _ in range(n_cols)]
    for (orig_r, orig_c), oo in origin_data.items():
        if orig_r >= header_row_count:
            continue
        for cc in range(orig_c, orig_c + oo['colspan']):
            if 0 <= cc < n_cols:
                header_texts_by_col[cc].append(oo['text'])
    for c in range(n_cols):
        col_values = []
        for r in range(header_row_count, n_rows):
            o = origin_data.get((r, c))
            if o is not None and o['rowspan'] == 1 and o['colspan'] == 1:
                col_values.append(o['text'])
        col_alignments.append(
            _infer_one_column_alignment(col_values, header_texts_by_col[c])
        )

    # Build tbl
    tbl = _make('tbl')
    tblPr = _make('tblPr')
    tblPr.append(_make('tblW', w=total_width, type='dxa'))
    tblPr.append(_make('tblInd', w=table_indent, type='dxa'))
    borders = _make('tblBorders')
    for bn in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        borders.append(_border_el(bn))
    tblPr.append(borders)
    tblPr.append(_make('tblLayout', type='fixed'))
    tbl.append(tblPr)

    tblGrid = _make('tblGrid')
    for w in column_widths:
        tblGrid.append(_make('gridCol', w=w))
    tbl.append(tblGrid)

    # Emit rows
    for r in range(n_rows):
        tr = _make('tr')
        c = 0
        while c < n_cols:
            st = status[r][c]
            in_header = (r < header_row_count)
            if st is None:
                # No origin, no cover — render a blank single cell
                tr.append(_make_merged_cell(
                    '', col_widths=[column_widths[c]], colspan=1,
                    is_header=in_header, bold=in_header,
                    alignment='center' if in_header else 'left',
                    font_size_half_pt=font_size_half_pt, vmerge=None))
                c += 1
            elif st[0] == 'ORIGIN':
                _, rs, cs = st
                o = origin_data[(r, c)]
                bold = o['bold'] if o['bold'] is not None else in_header
                if o['align']:
                    align = o['align']
                elif in_header:
                    align = 'center'
                elif rs > 1 or cs > 1:
                    # Merged body cells: center for readability (a spanned
                    # label looks cramped left-aligned in a wide cell).
                    align = 'center'
                else:
                    align = col_alignments[c]
                tr.append(_make_merged_cell(
                    o['text'],
                    col_widths=column_widths[c:c + cs],
                    colspan=cs, is_header=in_header, bold=bold,
                    alignment=align,
                    font_size_half_pt=font_size_half_pt,
                    vmerge='restart' if rs > 1 else None))
                c += cs
            elif st[0] == 'HCOVER':
                # Already covered by same-row origin's gridSpan
                c += 1
            elif st[0] == 'VCOVER_LEFT':
                _, cs = st
                tr.append(_make_merged_cell(
                    '', col_widths=column_widths[c:c + cs],
                    colspan=cs, is_header=in_header, bold=False,
                    alignment='center',
                    font_size_half_pt=font_size_half_pt,
                    vmerge='continue'))
                c += cs
            else:  # VCOVER_MID — covered by the VCOVER_LEFT cell's gridSpan
                c += 1
        tbl.append(tr)

    return tbl


def render_table_element(b: dict) -> OxmlElement:
    if isinstance(b.get('cells'), list) and len(b['cells']) > 0:
        return _render_cells_table(b)
    cols = b.get('columns') or []
    rows = b.get('rows') or []
    font_pt = _clamp_table_font_pt(b.get('fontSize', 12))
    font_size_half_pt = round(font_pt * 2)
    alignments_in = b.get('alignments')
    if isinstance(alignments_in, list) and len(alignments_in) == len(cols):
        alignments = alignments_in
    else:
        alignments = infer_column_alignments(cols, rows)
    table_indent = measure_dxa(PREFIX['dash'])
    total_width = PAGE_CONTENT_WIDTH - table_indent
    column_widths = (resolve_column_widths(b, cols, rows, total_width, font_pt)
                     if cols else [])

    tbl = _make('tbl')
    tblPr = _make('tblPr')
    tblPr.append(_make('tblW', w=total_width, type='dxa'))
    tblPr.append(_make('tblInd', w=table_indent, type='dxa'))
    borders = _make('tblBorders')
    for b_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        borders.append(_border_el(b_name))
    tblPr.append(borders)
    tblPr.append(_make('tblLayout', type='fixed'))
    tbl.append(tblPr)

    # tblGrid
    tblGrid = _make('tblGrid')
    for w in column_widths:
        tblGrid.append(_make('gridCol', w=w))
    tbl.append(tblGrid)

    # Header row
    tr = _make('tr')
    for idx, h in enumerate(cols):
        tr.append(_make_cell(h, is_header=True, col_width=column_widths[idx],
                             font_size_half_pt=font_size_half_pt,
                             alignment=alignments[idx] if idx < len(alignments) else 'left'))
    tbl.append(tr)
    # Body rows
    for row in rows:
        tr = _make('tr')
        for idx, _ in enumerate(cols):
            cell_val = row[idx] if idx < len(row) else ''
            tr.append(_make_cell(cell_val, is_header=False, col_width=column_widths[idx],
                                 font_size_half_pt=font_size_half_pt,
                                 alignment=alignments[idx] if idx < len(alignments) else 'left'))
        tbl.append(tr)
    return tbl


# ---------- Bullet rendering ----------------------------------
def render_bullet(doc: Document, bullet: dict, is_last: bool = False):
    kind = bullet.get('kind')

    if kind == 'image':
        render_image_bullet(doc, bullet, is_last=is_last)
        return

    if kind == 'table':
        caption = bullet.get('caption')
        unit = bullet.get('unit')
        if caption:
            cap_p = body_paragraph(f"< {caption} >", alignment='center',
                                   after=0, size=SIZE_BODY, bold=True)
            _apply_keep_next(cap_p)
            _append_block(doc, cap_p)
        if unit:
            unit_p = body_paragraph(f"({unit})", alignment='right',
                                    after=0, size=20)  # 10pt
            _apply_keep_next(unit_p)
            _append_block(doc, unit_p)
        tbl = render_table_element(bullet)
        _apply_table_keep_together(tbl)
        _append_block(doc, tbl)
        # Trailing spacer paragraph (matches JS: tiny line=40 exact, carries after-spacing)
        sp = make_paragraph([make_run('')], line=40, line_rule='exact', before=0,
                            after=(AFTER['last_in_section'] if is_last else AFTER['body']))
        _append_block(doc, sp)
        return

    if kind == 'subhead':
        p = body_paragraph(f"< {bullet.get('text', '')} >",
                           after=AFTER['body'], alignment='center', bold=True)
        _append_block(doc, p)
        return

    # Compute prefix
    if kind == 'circled':
        idx = (bullet.get('index') or 1) - 1
        circle = CIRCLED[idx] if 0 <= idx < len(CIRCLED) else CIRCLED[0]
        prefix = f' {circle} '
    elif kind == 'ordered':
        prefix = f" {bullet.get('index') or 1}) "
    else:
        prefix = PREFIX.get(kind, PREFIX['dash'])

    ann = normalize_annotation(bullet.get('annotation'))
    is_bold = bullet.get('bold') is True
    user_scale = bullet.get('scale')
    manual_scale = user_scale is not None
    prefix_scale = user_scale if manual_scale else WIDTH_PCT
    text_scale = user_scale if manual_scale else WIDTH_PCT
    text = bullet.get('text') or ''

    last_normal_after = AFTER['last_in_section'] if is_last else AFTER['body']

    hang_indent = hanging_indent_for(prefix, text)
    continuation_indent = hang_indent['left'] if hang_indent else measure_dxa(prefix)
    first_line_capacity = round(PAGE_CONTENT_WIDTH * WRAP_SAFETY_MARGIN)
    continuation_capacity = round((PAGE_CONTENT_WIDTH - continuation_indent) * WRAP_SAFETY_MARGIN)
    lines = split_bullet_text(prefix, text,
                              first_line_capacity=first_line_capacity,
                              continuation_capacity=continuation_capacity,
                              prefix_scale=prefix_scale,
                              text_scale=text_scale)

    # Auto widow fix: only when user didn't set `scale` manually. Shrink the
    # body's 장평 (prefix stays at WIDTH_PCT) in steps 94→90 and pick the first
    # candidate that eliminates the ≤2-word tail. If 90% still cannot resolve
    # the widow, revert to the original wrap + default scale.
    if not manual_scale and _last_line_is_widow(lines):
        for candidate in (94, 93, 92, 91, 90):
            cand_lines = split_bullet_text(
                prefix, text,
                first_line_capacity=first_line_capacity,
                continuation_capacity=continuation_capacity,
                prefix_scale=WIDTH_PCT,
                text_scale=candidate,
            )
            if not _last_line_is_widow(cand_lines):
                lines = cand_lines
                text_scale = candidate
                break
        # else: loop fell through; keep original `lines` and text_scale=WIDTH_PCT

    anchor_line_idx = -1
    if ann and ann['anchor']:
        for i, ln in enumerate(lines):
            if ann['anchor'] in ln:
                anchor_line_idx = i
                break

    use_anchor_placement = ann is not None and anchor_line_idx >= 0
    ann_insert_after = (anchor_line_idx if use_anchor_placement
                        else (len(lines) - 1 if ann else -1))

    ann_indent_dxa = 0
    if ann and ann['anchor']:
        if anchor_line_idx >= 0:
            lt = lines[anchor_line_idx]
            pos = lt.find(ann['anchor'])
            if anchor_line_idx == 0:
                ann_indent_dxa = (measure_dxa(prefix, prefix_scale)
                                  + measure_dxa(lt[:pos], text_scale))
            else:
                ann_indent_dxa = continuation_indent + measure_dxa(lt[:pos], text_scale)
        else:
            ann_indent_dxa = compute_anchor_indent_dxa(prefix, text, ann['anchor'])

    # Continuation-line leading spaces: match the first-line prefix width using
    # repeated half-width spaces so the continuation's x-start is reached purely
    # by literal whitespace (no paragraph w:ind). All default prefixes are an
    # exact multiple of half-width space width; under widow-fix (text_scale <
    # prefix_scale) the alignment may drift by ≤1 space, which is acceptable
    # per spec.
    cont_space_dxa = round(BODY_PT * 20 * 0.50 * (text_scale / 100))
    cont_spaces = (round(continuation_indent / cont_space_dxa)
                   if cont_space_dxa > 0 else 0)
    cont_pad = ' ' * max(cont_spaces, 0)

    for i, line_text in enumerate(lines):
        is_last_line = (i == len(lines) - 1)
        indent = hang_indent if i == 0 else None

        if is_last_line:
            after = last_normal_after
        else:
            after = AFTER['body_wrap']

        # First line carries the prefix at prefix_scale, body text at text_scale
        # (two runs); continuation lines pad with half-width spaces to align
        # under the first-line text start, without using paragraph indent.
        if i == 0:
            runs = [
                make_run(prefix, size=SIZE_BODY, bold=is_bold, scale=prefix_scale),
                make_run(line_text, size=SIZE_BODY, bold=is_bold, scale=text_scale),
            ]
        else:
            runs = [make_run(cont_pad + line_text, size=SIZE_BODY, bold=is_bold, scale=text_scale)]
        p = make_paragraph(runs, alignment=None, line=240, line_rule='auto',
                           after=after, indent=indent)
        _append_block(doc, p)

        if ann and i == ann_insert_after:
            # DrawingML 텍스트 상자 (텍스트 앞 wrap) — the owning bullet carries
            # its normal trailing after-spacing and the drawing floats without
            # reserving vertical space, so the host paragraph's after stays 0.
            _append_block(doc, annotation_paragraph(
                ann['text'],
                indent_dxa=ann_indent_dxa,
                after=0,
            ))


# ---------- Section heading & section -------------------------
def render_section_heading(doc: Document, section: dict):
    has_numeral = section.get('numeral') is not None and str(section.get('numeral')).strip() != ''
    has_number = (not has_numeral
                  and section.get('number') is not None
                  and str(section.get('number')).strip() != '')
    if has_numeral:
        prefix = f"{str(section['numeral']).strip()}. "
        heading_size = 32
    elif has_number:
        prefix = f"{str(section['number']).strip()}. "
        heading_size = 30
    else:
        prefix = '\u25A1 '
        heading_size = SIZE_BODY

    heading_text = f"{prefix}{section.get('heading', '')}"
    ann = normalize_annotation(section.get('headingAnnotation'))

    p = body_paragraph(heading_text,
                       after=AFTER['section_heading'],
                       indent=hanging_indent_for(prefix, section.get('heading', '')),
                       bold=True, size=heading_size)
    _append_block(doc, p)

    if ann:
        indent_dxa = compute_anchor_indent_dxa('', heading_text, ann['anchor'])
        _append_block(doc, annotation_paragraph(ann['text'],
                                                indent_dxa=indent_dxa,
                                                after=AFTER['section_heading']))


def render_section(doc: Document, section: dict, is_last: bool = False):
    render_section_heading(doc, section)
    bullets = section.get('bullets') or []
    for i, b in enumerate(bullets):
        render_bullet(doc, b, is_last=(i == len(bullets) - 1))
    if not is_last:
        sep = make_paragraph([make_run('', bold=True)],
                             line=240, line_rule='auto',
                             after=AFTER['section_separator'])
        _append_block(doc, sep)


# ---------- Title block, cover, TOC ---------------------------
def render_title_block(doc: Document, doc_data: dict):
    title = doc_data.get('title') or '(제목 없음)'
    p = make_paragraph([make_run(title, size=SIZE_TITLE, bold=True, underline=True)],
                       alignment='center', line=240, line_rule='auto', after=AFTER['title'])
    _append_block(doc, p)
    if doc_data.get('date'):
        p2 = make_paragraph([make_run(doc_data['date'])],
                            alignment='right', line=240, line_rule='auto', after=AFTER['date'])
        _append_block(doc, p2)


def _spacer_paragraph(size_half_pt: int = 24, after: int = 120) -> OxmlElement:
    return make_paragraph([make_run('', size=size_half_pt)],
                          line=240, line_rule='auto', after=after)


def _page_break_paragraph() -> OxmlElement:
    """A paragraph containing a single w:br (page break), matching JS PageBreak."""
    r = _make('r')
    r.append(_make('br', type='page'))
    return make_paragraph([r], line=240, line_rule='auto', after=0)


def _boxed_table(inner_children, *, width: int, alignment: str = 'center',
                 border_sz: int = 8, margins=(480, 480, 300, 300)) -> OxmlElement:
    """Single-cell bordered table used for TOC title and TOC content boxes."""
    tbl = _make('tbl')
    tblPr = _make('tblPr')
    tblPr.append(_make('tblW', w=width, type='dxa'))
    if alignment == 'center':
        tblPr.append(_make('jc', val='center'))
    elif alignment == 'right':
        tblPr.append(_make('jc', val='right'))
    borders = _make('tblBorders')
    for b_name in ('top', 'left', 'bottom', 'right'):
        borders.append(_make(b_name, val='single', sz=border_sz, color='000000', space=0))
    tblPr.append(borders)
    tblPr.append(_make('tblLayout', type='fixed'))
    tbl.append(tblPr)
    grid = _make('tblGrid')
    grid.append(_make('gridCol', w=width))
    tbl.append(grid)
    tr = _make('tr')
    tc = _make('tc')
    tcPr = _make('tcPr')
    tcPr.append(_make('tcW', w=width, type='dxa'))
    tc_borders = _make('tcBorders')
    for b_name in ('top', 'left', 'bottom', 'right'):
        tc_borders.append(_make(b_name, val='single', sz=border_sz, color='000000', space=0))
    tcPr.append(tc_borders)
    mar = _make('tcMar')
    top_m, bottom_m, left_m, right_m = margins
    mar.append(_make('top', w=top_m, type='dxa'))
    mar.append(_make('bottom', w=bottom_m, type='dxa'))
    mar.append(_make('left', w=left_m, type='dxa'))
    mar.append(_make('right', w=right_m, type='dxa'))
    tcPr.append(mar)
    tc.append(tcPr)
    for child in inner_children:
        tc.append(child)
    tr.append(tc)
    tbl.append(tr)
    return tbl


def render_toc_page(doc: Document, toc: dict):
    toc = toc or {}

    # 1) Small top spacer
    _append_block(doc, _spacer_paragraph(24, 120))

    # 2) Title box
    if toc.get('reportTitle'):
        box_width = 7600
        title_p = make_paragraph(
            [make_run(toc['reportTitle'], size=44, bold=True, scale=100)],
            alignment='center', line=240, line_rule='auto', after=0,
        )
        _append_block(doc, _boxed_table([title_p], width=box_width,
                                        alignment='center', border_sz=8,
                                        margins=(480, 480, 300, 300)))

    # 3) 2 empty spacers
    _append_block(doc, _spacer_paragraph(24, 120))
    _append_block(doc, _spacer_paragraph(24, 120))

    # 4) "- 목    차 -" center 18pt bold
    _append_block(doc, make_paragraph(
        [make_run('- 목    차 -', size=36, bold=True, scale=100)],
        alignment='center', line=240, line_rule='auto', after=300,
    ))

    # 5) TOC body box
    toc_children = []
    main_items = toc.get('items') or []
    for idx, item in enumerate(main_items):
        text = item if isinstance(item, str) else (item.get('text') or '')
        num = (item.get('number') if isinstance(item, dict) and item.get('number') is not None
               else (idx + 1))
        item_text = f"{num}. {text}"
        toc_children.append(make_paragraph(
            [make_run(item_text, size=32, bold=True, scale=100)],
            alignment='left', line=360, line_rule='auto', after=280,
            indent={'left': 600, 'hanging': 280},
        ))

    if toc.get('appendix'):
        label = toc['appendix'].get('label') or '[별첨]'
        app_items = toc['appendix'].get('items') or []
        for idx, item in enumerate(app_items):
            text = item if isinstance(item, str) else (item.get('text') or '')
            prefix = f"{label} " if idx == 0 else ''
            left_indent = 900 if idx == 0 else 1900
            toc_children.append(make_paragraph(
                [make_run(f"{prefix}{text}", size=28, scale=100)],
                alignment='left', line=300, line_rule='auto', after=200,
                indent={'left': left_indent},
            ))

    if toc_children:
        box_width = 8800
        _append_block(doc, _boxed_table(toc_children, width=box_width,
                                        alignment='center', border_sz=4,
                                        margins=(500, 500, 500, 500)))

    # Spacer
    _append_block(doc, _spacer_paragraph(24, 240))

    # 6) Date
    if toc.get('date'):
        _append_block(doc, make_paragraph(
            [make_run(toc['date'], size=36, scale=100)],
            alignment='center', line=240, line_rule='auto', after=160,
        ))

    # Small gap
    _append_block(doc, _spacer_paragraph(16, 80))

    # 7) Author
    if toc.get('author'):
        _append_block(doc, make_paragraph(
            [make_run(toc['author'], size=40, bold=True, scale=100,
                      character_spacing=58)],
            alignment='center', line=240, line_rule='auto', after=0,
        ))

    # 8) Page break
    _append_block(doc, _page_break_paragraph())


def render_cover(doc: Document, cover: dict):
    cover = cover or {}
    render_title_block(doc, {'title': cover.get('title') or '',
                             'date': cover.get('date') or ''})

    if cover.get('subtitle'):
        _append_block(doc, body_paragraph(cover['subtitle'],
                                          alignment='center', after=AFTER['body']))

    meta = cover.get('metaLines') or []
    for idx, line in enumerate(meta):
        last = (idx == len(meta) - 1)
        p = make_paragraph([make_run(line)], alignment='right',
                           line=240, line_rule='auto',
                           after=(AFTER['section_separator'] if last else 120))
        _append_block(doc, p)

    if cover.get('tocHeading'):
        render_section_heading(doc, {'heading': cover['tocHeading']})

    toc_items = cover.get('tocItems') or []
    notes = cover.get('notes') or []
    has_trailing_notes = len(notes) > 0
    for idx, item in enumerate(toc_items):
        render_bullet(doc, {'kind': 'dash', 'text': item},
                      is_last=(idx == len(toc_items) - 1 and not has_trailing_notes))

    for idx, note in enumerate(notes):
        render_bullet(doc, {'kind': 'note', 'text': note},
                      is_last=(idx == len(notes) - 1))

    _append_block(doc, _page_break_paragraph())


def render_closing(doc: Document):
    p = make_paragraph([make_run('- 이  상 -')], alignment='right',
                       line=240, line_rule='auto', after=AFTER['closing'])
    _append_block(doc, p)


def render_appendix_label(doc: Document, n: int, title: str | None):
    text = f"#별첨{n}. {title}" if title else f"#별첨{n}."
    # Paragraph combines a page break followed by the label run (matches JS layout).
    br_r = _make('r')
    br_r.append(_make('br', type='page'))
    label_r = make_run(text, size=SIZE_APPEND, bold=True)
    p = make_paragraph([br_r, label_r], line=240, line_rule='auto',
                       after=AFTER['append_label'])
    _append_block(doc, p)


def render_body(doc: Document, doc_data: dict):
    default_skip = bool(doc_data.get('toc'))
    skip_explicit = doc_data.get('skipTitleBlock')
    skip = default_skip if skip_explicit is None else bool(skip_explicit)
    if not skip:
        render_title_block(doc, doc_data)
    sections = doc_data.get('sections') or []
    for i, s in enumerate(sections):
        render_section(doc, s, is_last=(i == len(sections) - 1))
    render_closing(doc)


def render_appendix(doc: Document, app: dict, n: int):
    render_appendix_label(doc, n, app.get('title'))
    sections = app.get('sections') or []
    for i, s in enumerate(sections):
        render_section(doc, s, is_last=(i == len(sections) - 1))


# ---------- Post-processing: autoSpaceDE/DN -------------------
_AUTO_SPACE_SNIPPET = '<w:autoSpaceDE w:val="0"/><w:autoSpaceDN w:val="0"/>'
_PPR_OPEN_RE = re.compile(r'<w:pPr>')


def post_process_autospace(buf: bytes) -> bytes:
    """Inject <w:autoSpaceDE/> + <w:autoSpaceDN/> at the start of every <w:pPr>.

    Matches the JS postProcessAutoSpace(). We also need to handle paragraphs
    that don't yet have a pPr — python-docx omits pPr when no properties are
    set. To keep parity we ensure every paragraph has at least autoSpace
    entries by inserting pPr into paragraph marks without one.
    """
    with io.BytesIO(buf) as src_stream:
        with zipfile.ZipFile(src_stream, 'r') as zin:
            out_bytes = io.BytesIO()
            with zipfile.ZipFile(out_bytes, 'w', zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)
                    if item.filename == 'word/document.xml':
                        xml = data.decode('utf-8')
                        # 1) Prepend autoSpace into every existing <w:pPr>
                        xml = _PPR_OPEN_RE.sub('<w:pPr>' + _AUTO_SPACE_SNIPPET, xml)
                        # 2) Insert a pPr (with autoSpace only) into paragraphs
                        #    that don't already have one.  <w:p> → <w:p><w:pPr>…</w:pPr>
                        xml = re.sub(
                            r'<w:p(\s[^>]*)?>(?!<w:pPr>)',
                            lambda m: f'<w:p{m.group(1) or ""}><w:pPr>{_AUTO_SPACE_SNIPPET}</w:pPr>',
                            xml,
                        )
                        data = xml.encode('utf-8')
                    zout.writestr(item, data)
            return out_bytes.getvalue()


# ---------- Document setup & main -----------------------------
def _setup_page_numbers(doc: Document):
    """Center-aligned '- N -' page number in the footer at 12pt 바탕체.

    Footer distance itself is set in _setup_document_defaults via PAGE['footer']
    (510 twips ≈ 0.9cm).
    """
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    ftr_element = footer._element
    for p in ftr_element.findall(qn('w:p')):
        ftr_element.remove(p)

    p = _make('p')
    pPr = _make('pPr')
    pPr.append(_make('jc', val='center'))
    p.append(pPr)

    p.append(make_run('- ', size=SIZE_TABLE))

    fld = _make('fldSimple', instr='PAGE')
    fld.append(make_run('1', size=SIZE_TABLE))
    p.append(fld)

    p.append(make_run(' -', size=SIZE_TABLE))

    ftr_element.append(p)


def _setup_document_defaults(doc: Document):
    section = doc.sections[0]
    section.page_width = Twips(PAGE['width'])
    section.page_height = Twips(PAGE['height'])
    section.top_margin = Twips(PAGE['top'])
    section.right_margin = Twips(PAGE['right'])
    section.bottom_margin = Twips(PAGE['bottom'])
    section.left_margin = Twips(PAGE['left'])
    section.header_distance = Twips(PAGE['header'])
    section.footer_distance = Twips(PAGE['footer'])
    _setup_page_numbers(doc)

    # Set default run font on styles.xml → docDefaults/rPrDefault/rPr/rFonts
    styles_el = doc.styles.element
    rpr_defaults = styles_el.find(qn('w:docDefaults'))
    if rpr_defaults is not None:
        rpr_default = rpr_defaults.find(qn('w:rPrDefault'))
        if rpr_default is not None:
            rpr = rpr_default.find(qn('w:rPr'))
            if rpr is None:
                rpr = _make('rPr')
                rpr_default.append(rpr)
            # Clear any existing rFonts / sz to replace with our defaults
            for tag in ('rFonts', 'sz', 'szCs'):
                for el in rpr.findall(qn('w:' + tag)):
                    rpr.remove(el)
            _set_run_font(rpr)
            rpr.append(_make('sz', val=SIZE_BODY))
            rpr.append(_make('szCs', val=SIZE_BODY))

    # Remove the default empty paragraph that python-docx inserts.
    body = doc.element.body
    for p in body.findall(qn('w:p')):
        # Only remove the leading empty paragraph (no runs)
        if len(p.findall(qn('w:r'))) == 0:
            body.remove(p)
            break


def main():
    if len(sys.argv) < 3:
        sys.stderr.write('Usage: python generate.py <content.json> <output.docx>\n')
        sys.exit(1)
    content_path = sys.argv[1]
    out_path = sys.argv[2]

    with open(content_path, 'r', encoding='utf-8') as f:
        doc_data = json.load(f)

    doc = Document()
    _setup_document_defaults(doc)

    if doc_data.get('toc'):
        render_toc_page(doc, doc_data['toc'])
    if doc_data.get('cover'):
        render_cover(doc, doc_data['cover'])
    render_body(doc, doc_data)
    for i, app in enumerate(doc_data.get('appendices') or []):
        render_appendix(doc, app, i + 1)

    # Serialize to bytes, then post-process autoSpaceDE/DN
    buffer = io.BytesIO()
    doc.save(buffer)
    buf = buffer.getvalue()
    buf = post_process_autospace(buf)
    with open(out_path, 'wb') as f:
        f.write(buf)
    # Normalize path to forward slashes for readability.
    print(f"Wrote {out_path.replace(os.sep, '/')} ({len(buf)} bytes)")


if __name__ == '__main__':
    main()